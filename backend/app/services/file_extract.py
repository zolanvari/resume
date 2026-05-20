"""Résumé text extraction across every supported upload format.

PDF, .txt, modern .docx and legacy binary .doc all collapse to plain text
here. An LLM downstream normalises whatever layout survives, so extraction is
deliberately best-effort.
"""

import io
import os
import subprocess
import tempfile

from docx import Document

from app.services.pdf_parser import extract_text_from_pdf

_DOCX_CT = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
_DOC_CT = "application/msword"


class UnsupportedFileType(Exception):
    """Raised when an uploaded file is not a format we can read."""


def _extract_docx(data: bytes) -> str:
    """Modern Word (.docx) via python-docx - paragraphs and table cells."""
    doc = Document(io.BytesIO(data))
    parts: list[str] = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append("  ".join(cells))
    return "\n".join(parts)


def _extract_doc(data: bytes) -> str:
    """Legacy binary Word (.doc) via the `antiword` CLI (installed in the image)."""
    with tempfile.NamedTemporaryFile(suffix=".doc", delete=False) as tmp:
        tmp.write(data)
        tmp_path = tmp.name
    try:
        result = subprocess.run(
            ["antiword", tmp_path],
            capture_output=True,
            timeout=20,
        )
    except FileNotFoundError as exc:  # antiword missing from the image
        raise UnsupportedFileType(
            "Legacy .doc support is unavailable. Save the file as .docx or PDF."
        ) from exc
    except subprocess.TimeoutExpired as exc:
        raise UnsupportedFileType("The .doc file took too long to read.") from exc
    finally:
        try:
            os.unlink(tmp_path)
        except OSError:
            pass
    if result.returncode != 0:
        raise UnsupportedFileType(
            "Could not read this .doc file. Save it as .docx or PDF and retry."
        )
    return result.stdout.decode("utf-8", errors="replace")


def extract_resume_text(filename: str, content_type: str, data: bytes) -> str:
    """Dispatch on filename/content-type and return extracted plain text.

    Raises UnsupportedFileType for formats we cannot read.
    """
    name = (filename or "").lower()
    ctype = (content_type or "").lower()

    if name.endswith(".pdf") or "pdf" in ctype:
        return extract_text_from_pdf(data)
    if name.endswith(".docx") or ctype == _DOCX_CT:
        return _extract_docx(data)
    if name.endswith(".doc") or ctype == _DOC_CT:
        return _extract_doc(data)
    if name.endswith(".txt") or ctype.startswith("text/"):
        return data.decode("utf-8", errors="replace")

    raise UnsupportedFileType(
        "Unsupported file type. Upload a PDF, Word (.doc/.docx), or .txt file, "
        "or paste your résumé as text."
    )
